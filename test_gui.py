import tkinter as tk
from tkinter import messagebox, simpledialog, Toplevel, ttk, filedialog
from docx import Document
import random
import time
from PIL import Image, ImageTk, ImageDraw, ImageGrab
import os
import tempfile
import atexit
from lxml import etree
import json
import sys
from tkinter import font as tkfont
import glob
import io
import docx
import requests
import ctypes
import openpyxl
from datetime import datetime
import traceback  # Qo'shildi

# API URL
API_BASE_URL = "https://test-system-api-wggo.onrender.com"

def api_login(login, password):
    """API orqali login qilish"""
    try:
        response = requests.post(f"{API_BASE_URL}/auth/login", 
                               json={"login": login, "password": password})
        if response.status_code == 200:
            return response.json()
        return None
    except Exception as e:
        print(f"Login xatosi: {str(e)}")
        return None

def api_get_questions():
    """API dan savollarni olish"""
    try:
        response = requests.get(f"{API_BASE_URL}/questions")
        if response.status_code == 200:
            return response.json()
        return None
    except Exception as e:
        print(f"Savollarni olishda xato: {str(e)}")
        return None

def api_save_result(student_name, score, total_questions, answers):
    """Test natijasini API ga yuborish"""
    try:
        data = {
            "student_name": student_name,
            "score": score,
            "total_questions": total_questions,
            "answers": answers,
            "date": datetime.now().isoformat()
        }
        response = requests.post(f"{API_BASE_URL}/results", json=data)
        return response.status_code == 200
    except Exception as e:
        print(f"Natijani saqlashda xato: {str(e)}")
        return False

# Telegram bot token (siz bergan token)
TELEGRAM_BOT_TOKEN = "8061838746:AAFAstWBE4ROl_y5lj4lnldKg1v5oxKv4kM"
# Admin chat_id ni dastlab None qilib qo'yamiz, birinchi xabarda aniqlaymiz
global_admin_chat_id = None

# PyInstaller uchun universal pathlar
if getattr(sys, 'frozen', False):
    # PyInstaller exe uchun
    BASE_DIR = os.path.dirname(sys.executable)
    RESOURCE_DIR = os.path.join(sys._MEIPASS, 'resources')
else:
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))
    RESOURCE_DIR = os.path.join(BASE_DIR, 'resources')

# 2. STUDENTS_FILE ni universal qilaman:
STUDENTS_FILE = os.path.join(BASE_DIR, 'students.xlsx')

def ensure_students_file():
    if not os.path.exists(STUDENTS_FILE):
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = 'Students'
        ws.append(['ISM', 'FAMILIYA', 'LOGIN', 'PAROL'])
        wb.save(STUDENTS_FILE)

def read_students_list():
    ensure_students_file()
    wb = openpyxl.load_workbook(STUDENTS_FILE)
    ws = wb.active
    return [
        {
            'ism': str(row[0].value).strip().upper(),
            'familiya': str(row[1].value).strip().upper(),
            'login': str(row[2].value).strip(),
            'parol': str(row[3].value).strip()
        }
        for row in ws.iter_rows(min_row=2) if all(cell.value for cell in row[:4])
    ]

def add_student_to_list(ism, familiya, login, parol):
    ensure_students_file()
    wb = openpyxl.load_workbook(STUDENTS_FILE)
    ws = wb.active
    ws.append([ism.strip().upper(), familiya.strip().upper(), login.strip(), parol.strip()])
    wb.save(STUDENTS_FILE)

def remove_student_from_list(login):
    ensure_students_file()
    wb = openpyxl.load_workbook(STUDENTS_FILE)
    ws = wb.active
    login = login.strip()
    rows = list(ws.iter_rows(min_row=2))
    for i, row in enumerate(rows, start=2):
        if str(row[2].value).strip() == login:
            ws.delete_rows(i)
            break
    wb.save(STUDENTS_FILE)

def send_result_to_telegram(result_text, bot_token, chat_id):
    url = f"https://api.telegram.org/bot{bot_token}/sendMessage"
    data = {
        "chat_id": chat_id,
        "text": result_text,
        "parse_mode": "HTML"  # HTML formatini qo'shdik
    }
    try:
        print(f"\nTelegram API so'rovi:\nURL: {url}\nData: {data}")
        response = requests.post(url, data=data, timeout=10)  # Timeout qo'shildi
        print(f"Telegram javob kodi: {response.status_code}")
        print(f"Telegram javob matni: {response.text}")
        
        if response.status_code != 200:
            print(f"Telegram xatosi: {response.text}")
            return False
        return True
    except requests.Timeout:
        print("Telegram so'rovi vaqt chegarasidan oshib ketdi!")
        return False
    except requests.RequestException as e:
        print(f"Telegram so'rovida xato: {str(e)}")
        return False
    except Exception as e:
        print(f"Telegramga yuborishda xato: {str(e)}")
        traceback.print_exc()  # To'liq xato stacktrace'ni chiqarish
        return False

def get_latest_chat_id(bot_token):
    url = f"https://api.telegram.org/bot{bot_token}/getUpdates"
    try:
        print(f"\nChat ID so'rovi:\nURL: {url}")
        resp = requests.get(url, timeout=5)
        print(f"Chat ID so'rovi javob kodi: {resp.status_code}")
        print(f"Chat ID so'rovi javob matni: {resp.text}")
        
        data = resp.json()
        if 'result' in data and data['result']:
            # Eng oxirgi xabar yuborgan chat_id ni olamiz
            for update in reversed(data['result']):
                if 'message' in update and 'chat' in update['message']:
                    chat_id = str(update['message']['chat']['id'])
                    print(f"Topilgan chat ID: {chat_id}")
                    return chat_id
        print("Chat ID topilmadi!")
        return None
    except Exception as e:
        print(f"Chat ID olishda xato: {str(e)}")
        traceback.print_exc()
        return None

def send_result_screenshot_to_telegram(bot_token, chat_id, image_path, caption=None):
    url = f"https://api.telegram.org/bot{bot_token}/sendPhoto"
    data = {"chat_id": chat_id}
    if caption:
        data["caption"] = caption
    try:
        with open(image_path, "rb") as img_file:
            files = {"photo": img_file}
            requests.post(url, data=data, files=files)
    except Exception as e:
        print(f"Telegramga rasm yuborishda xato: {e}")

def get_window_bbox(root):
    root.update()
    x = root.winfo_rootx()
    y = root.winfo_rooty()
    w = root.winfo_width()
    h = root.winfo_height()
    return (x, y, x + w, y + h)

def get_widget_bbox(widget):
    widget.update()
    x = widget.winfo_rootx()
    y = widget.winfo_rooty()
    w = widget.winfo_width()
    h = widget.winfo_height()
    return (x, y, x + w, y + h)

class ModernTestApp:
    def __init__(self, root):
        self.root = root
        self.setup_window()
        self.setup_styles()
        self.initialize_variables()
        self.load_data()
        self.create_start_screen()
        # Xatolarni tahlil qilish uchun o'zgaruvchilar
        self.reviewed_questions = set()
        self.review_start_time = None

    def setup_window(self):
        self.root.title("Test Dasturi")
        self.root.geometry("1200x800")
        self.root.minsize(1000, 700)
        self.root.protocol("WM_DELETE_WINDOW", self.on_close)
        self.root.configure(bg="#f5f5f5")

    def setup_styles(self):
        self.style = ttk.Style()
        self.style.theme_use('clam')
        self.colors = {
            'primary': '#006225',
            'secondary': '#166088',
            'accent': '#4fc3f7',
            'light': '#f8f9fa',
            'dark': '#343a40',
            'success': '#28a745',
            'danger': '#dc3545',
            'warning': '#ffc107',
            'text': '#212529',
            'border': '#dee2e6'
        }
        self.style.configure('TFrame', background=self.colors['light'])
        self.style.configure('TLabel', background=self.colors['light'], font=('Segoe UI', 11))
        self.style.configure('TButton', font=('Segoe UI', 11), padding=8, borderwidth=1)
        self.style.configure('Primary.TButton', background=self.colors['primary'], foreground='white')
        self.style.configure('Secondary.TButton', background=self.colors['secondary'], foreground='white')
        self.style.configure('Accent.TButton', background=self.colors['accent'], foreground='white')
        self.style.configure('Success.TButton', background=self.colors['success'], foreground='white')
        self.style.configure('Danger.TButton', background=self.colors['danger'], foreground='white')
        self.style.configure('Warning.TButton', background=self.colors['warning'], foreground='black')
        self.style.configure('TEntry', font=('Segoe UI', 11), padding=6, bordercolor=self.colors['border'])
        self.style.configure('Card.TFrame', background='white', borderwidth=1, relief='solid', bordercolor=self.colors['border'])
        self.style.configure('Progress.Horizontal.TProgressbar', thickness=20, background=self.colors['primary'])
        self.style.map('TButton',
                      background=[('active', self.colors['primary']), ('disabled', self.colors['light'])],
                      bordercolor=[('active', self.colors['primary'])])

    def initialize_variables(self):
        self.username = ""
        self.all_questions = []
        self.current_test_set = []
        self.user_answers = {}
        self.current_question = 0
        self.timer_seconds = 45 * 60
        self.timer_running = False
        self.images = []
        self.temp_files = []
        self.admin_password = "0088"
        self.zoom_window = None
        self.admin_override = False
        self.results_shown = False
        self.current_zoom = 1.0
        self.used_question_indices = set()
        self.current_screen = ""
        self.settings = {
            "test_time": 45,
            "questions_per_test": 30,
            "shuffle_answers": True,
            "dark_mode": False
        }
        self.admin_image_path = None
        self.admin_image_tk = None

    def load_data(self):
        self.load_settings()
        self.load_used_questions()
        atexit.register(self.cleanup)

    def cleanup(self):
        self.cleanup_temp_files()
        self.save_used_questions()
        self.save_settings()

    def cleanup_temp_files(self):
        for filepath in self.temp_files:
            try:
                if os.path.exists(filepath):
                    os.remove(filepath)
            except Exception as e:
                print(f"Faylni o'chirishda xato: {e}")

    def save_used_questions(self):
        try:
            with open('used_questions.json', 'w') as f:
                json.dump(list(self.used_question_indices), f)
        except Exception as e:
            print(f"Savollarni saqlashda xato: {e}")

    def load_used_questions(self):
        try:
            if os.path.exists('used_questions.json'):
                with open('used_questions.json', 'r') as f:
                    data = json.load(f)
                    self.used_question_indices = set(data) if data else set()
        except Exception as e:
            print(f"Savollarni yuklashda xato: {e}")
            self.used_question_indices = set()

    def save_settings(self):
        try:
            with open('settings.json', 'w') as f:
                json.dump(self.settings, f)
        except Exception as e:
            print(f"Sozlamalarni saqlashda xato: {e}")

    def load_settings(self):
        try:
            if os.path.exists('settings.json'):
                with open('settings.json', 'r') as f:
                    loaded_settings = json.load(f)
                    self.settings.update(loaded_settings)
                    if self.settings['dark_mode']:
                        self.toggle_dark_mode(True)
        except Exception as e:
            print(f"Sozlamalarni yuklashda xato: {e}")

    def toggle_dark_mode(self, enabled=None):
        if enabled is None:
            self.settings['dark_mode'] = not self.settings['dark_mode']
        else:
            self.settings['dark_mode'] = enabled
        if self.settings['dark_mode']:
            self.colors.update({
                'light': '#212529',
                'dark': '#f8f9fa',
                'text': '#f8f9fa',
                'border': '#495057'
            })
        else:
            self.colors.update({
                'light': '#f8f9fa',
                'dark': '#212529',
                'text': '#212529',
                'border': '#dee2e6'
            })
        self.setup_styles()
        self.refresh_ui()

    def refresh_ui(self):
        if hasattr(self, 'current_screen'):
            if self.current_screen == "start":
                self.create_start_screen()
            elif self.current_screen == "test":
                self.show_test_screen()
            elif self.current_screen == "results":
                self.finish_test()

    def create_start_screen(self):
        self.current_screen = "start"
        self.clear_screen()
        self.root.geometry("1200x800")
        main_frame = ttk.Frame(self.root, padding=(40, 20))
        main_frame.pack(expand=True, fill='both')
        # LOGO
        try:
            from PIL import Image, ImageTk
            logo_path = os.path.join(RESOURCE_DIR, "logo.png")
            logo_img = Image.open(logo_path)
            logo_img = logo_img.resize((120, 120), Image.LANCZOS)
            self.logo_photo = ImageTk.PhotoImage(logo_img)
            logo_label = ttk.Label(main_frame, image=self.logo_photo)
            logo_label.pack(pady=(0, 10))
        except Exception as e:
            print(f"Logo yuklanmadi: {e}")
        header_frame = ttk.Frame(main_frame)
        header_frame.pack(pady=(0, 30))
        ttk.Label(
            header_frame, 
            text="Test Dasturiga Xush Kelibsiz!", 
            font=('Segoe UI', 24, 'bold'),
            foreground=self.colors['primary']
        ).pack()
        input_card = ttk.Frame(main_frame, style='Card.TFrame', padding=30)
        input_card.pack(pady=10, ipadx=20, ipady=20)
        ttk.Label(
            input_card, 
            text="Login:", 
            font=('Segoe UI', 14)
        ).pack(pady=(0, 5))
        self.login_entry = ttk.Entry(
            input_card, 
            font=('Segoe UI', 14), 
            width=30
        )
        self.login_entry.pack(pady=5, ipady=5)
        ttk.Label(
            input_card, 
            text="Parol:", 
            font=('Segoe UI', 14)
        ).pack(pady=(10, 5))
        self.password_entry = ttk.Entry(
            input_card, 
            font=('Segoe UI', 14), 
            width=30,
            show='*'
        )
        self.password_entry.pack(pady=5, ipady=5)
        self.login_entry.focus_set()
        # ICONLAR
        try:
            self.icon_play = ImageTk.PhotoImage(Image.open(os.path.join(RESOURCE_DIR, "play.png")).resize((32, 32), Image.LANCZOS))
            self.icon_settings = ImageTk.PhotoImage(Image.open(os.path.join(RESOURCE_DIR, "settings.png")).resize((32, 32), Image.LANCZOS))
            self.icon_exit = ImageTk.PhotoImage(Image.open(os.path.join(RESOURCE_DIR, "exit.png")).resize((32, 32), Image.LANCZOS))
        except Exception as e:
            print(f"Iconlar yuklanmadi: {e}")
            self.icon_play = self.icon_settings = self.icon_exit = None
        button_frame = ttk.Frame(main_frame)
        button_frame.pack(pady=20, fill='x')
        ttk.Button(
            button_frame, 
            text="Boshlash", 
            style='Primary.TButton',
            command=self.start_test,
            image=self.icon_play if hasattr(self, 'icon_play') else None,
            compound='left'
        ).pack(side='left', padx=5, pady=10, expand=True, fill='x')
        ttk.Button(
            button_frame,
            text="Admin Panel",
            style='Accent.TButton',
            command=self.show_admin_panel,
            image=self.icon_settings if hasattr(self, 'icon_settings') else None,
            compound='left'
        ).pack(side='left', padx=5, pady=10, expand=True, fill='x')
        ttk.Button(
            button_frame, 
            text="Chiqish", 
            style='Danger.TButton',
            command=self.on_close,
            image=self.icon_exit if hasattr(self, 'icon_exit') else None,
            compound='left'
        ).pack(side='left', padx=5, pady=10, expand=True, fill='x')
        self.login_entry.bind('<Return>', lambda event: self.password_entry.focus_set())
        self.password_entry.bind('<Return>', lambda event: self.start_test())

    def show_settings(self):
        settings_window = Toplevel(self.root)
        settings_window.title("Sozlamalar")
        settings_window.geometry("500x400")
        settings_window.resizable(False, False)
        main_frame = ttk.Frame(settings_window, padding=20)
        main_frame.pack(expand=True, fill='both')
        ttk.Label(
            main_frame, 
            text="Test vaqti (daqiqa):", 
            font=('Segoe UI', 12)
        ).pack(pady=(10, 0))
        time_entry = ttk.Entry(main_frame, font=('Segoe UI', 12))
        time_entry.insert(0, str(self.settings["test_time"]))
        time_entry.pack(fill='x', pady=5)
        ttk.Label(
            main_frame, 
            text="Savollar soni:", 
            font=('Segoe UI', 12)
        ).pack(pady=(10, 0))
        questions_entry = ttk.Entry(main_frame, font=('Segoe UI', 12))
        questions_entry.insert(0, str(self.settings["questions_per_test"]))
        questions_entry.pack(fill='x', pady=5)
        options_frame = ttk.Frame(main_frame)
        options_frame.pack(pady=10, fill='x')
        shuffle_var = tk.BooleanVar(value=self.settings["shuffle_answers"])
        ttk.Checkbutton(
            options_frame, 
            text="Javoblarni aralashtirish", 
            variable=shuffle_var
        ).pack(anchor='w')
        dark_mode_var = tk.BooleanVar(value=self.settings["dark_mode"])
        ttk.Checkbutton(
            options_frame, 
            text="Qorong'i rejim", 
            variable=dark_mode_var,
            command=lambda: self.toggle_dark_mode(dark_mode_var.get())
        ).pack(anchor='w', pady=(10, 0))
        ttk.Label(
            main_frame, 
            text="Admin parolini o'zgartirish:", 
            font=('Segoe UI', 12)
        ).pack(pady=(20, 5))
        pass_frame = ttk.Frame(main_frame)
        pass_frame.pack(fill='x', pady=5)
        ttk.Label(pass_frame, text="Joriy parol:").pack(side='left')
        current_pass = ttk.Entry(pass_frame, show='*', width=15)
        current_pass.pack(side='left', padx=5)
        ttk.Label(pass_frame, text="Yangi parol:").pack(side='left', padx=(10, 0))
        new_pass = ttk.Entry(pass_frame, show='*', width=15)
        new_pass.pack(side='left', padx=5)
        def save_settings():
            # Admin parolini so'rash
            password = simpledialog.askstring(
                "Admin paroli",
                "Sozlamalarni saqlash uchun admin parolini kiriting:",
                show='*',
                parent=settings_window
            )
            if password != self.admin_password:
                messagebox.showerror("Xato", "Noto'g'ri admin paroli! Sozlamalar saqlanmadi.")
                return
            try:
                new_time = int(time_entry.get())
                new_questions = int(questions_entry.get())
                if new_time <= 0 or new_questions <= 0:
                    raise ValueError("Qiymatlar musbat son bo'lishi kerak")
                self.settings.update({
                    "test_time": new_time,
                    "questions_per_test": new_questions,
                    "shuffle_answers": shuffle_var.get(),
                    "dark_mode": dark_mode_var.get()
                })
                if current_pass.get() and new_pass.get():
                    if current_pass.get() == self.admin_password:
                        self.admin_password = new_pass.get()
                        messagebox.showinfo("Muvaffaqiyat", "Parol muvaffaqiyatli o'zgartirildi!")
                    else:
                        messagebox.showerror("Xato", "Noto'g'ri joriy parol!")
                self.save_settings()
                settings_window.destroy()
                messagebox.showinfo("Muvaffaqiyat", "Sozlamalar saqlandi!")
            except ValueError as e:
                messagebox.showerror("Xato", f"Noto'g'ri qiymat: {e}")
        ttk.Button(
            main_frame, 
            text="Saqlash", 
            style='Primary.TButton',
            command=save_settings
        ).pack(pady=10, fill='x')

    def start_test(self):
        if not self.login_entry or not self.password_entry:
            messagebox.showwarning("Xato", "Login yoki parol maydoni topilmadi!")
            return
            
        login = self.login_entry.get().strip()
        parol = self.password_entry.get().strip()
        
        if not login or not parol:
            messagebox.showwarning("Ogohlantirish", "Iltimos, login va parolni kiriting!")
            return
            
        try:
            # API orqali login qilish
            api_result = api_login(login, parol)
            if not api_result:
                messagebox.showerror("Xato", "Login yoki parol noto'g'ri yoki serverga ulanishda xatolik!")
                return
                
            self.username = api_result.get('name', f"{login}")
            
            # API dan savollarni olish
            questions = api_get_questions()
            if not questions:
                messagebox.showerror("Xato", "Serverdan savollarni olishda xatolik!")
                return
                
            self.all_questions = questions
            if not self.all_questions:
                messagebox.showerror("Xato", "Test bazasida savollar topilmadi!")
                return
                
            if len(self.all_questions) < self.settings["questions_per_test"]:
                messagebox.showerror("Xato", 
                                   f"Testlar soni {self.settings['questions_per_test']} tadan kam! "
                                   f"Mavjud testlar soni: {len(self.all_questions)}")
                return
                
            # Timer ni qayta o'rnatish
            self.timer_seconds = self.settings["test_time"] * 60
            self.timer_running = True
            
            self.start_new_test()
            
        except Exception as e:
            messagebox.showerror("Xato", f"Test boshlashda xatolik: {str(e)}")
            return

    def load_all_tests_from_folder(self, folder_path):
        all_questions = []
        # Faqat tanlangan fayllarni yuklash
        enabled_files = None
        if os.path.exists('enabled_baza_files.json'):
            try:
                with open('enabled_baza_files.json', 'r') as f:
                    enabled_files = set(json.load(f))
            except Exception:
                enabled_files = None
        docx_files = glob.glob(os.path.join(folder_path, "*.docx"))
        for file in docx_files:
            if enabled_files is not None and os.path.basename(file) not in enabled_files:
                continue
            questions = self.load_tests_from_docx(file)
            for idx, q in enumerate(questions):
                q['source_file'] = os.path.basename(file)
                q['source_index'] = idx + 1
            all_questions.extend(questions)
        return all_questions

    def load_tests_from_docx(self, filepath):
        try:
            doc = Document(filepath)
            questions = []
            for table in doc.tables:
                for row in table.rows[1:]:
                    cells = row.cells
                    if len(cells) < 5:
                        continue
                    question_text = cells[0].text.strip()
                    question_images = self.extract_images_from_element(cells[0]._element, doc)
                    true_answer = cells[1].text.strip()
                    true_answer_images = self.extract_images_from_element(cells[1]._element, doc)
                    false_answers = [
                        {"text": cells[2].text.strip(), "images": self.extract_images_from_element(cells[2]._element, doc)},
                        {"text": cells[3].text.strip(), "images": self.extract_images_from_element(cells[3]._element, doc)},
                        {"text": cells[4].text.strip(), "images": self.extract_images_from_element(cells[4]._element, doc)}
                    ]
                    options_data = [
                        {"text": true_answer, "images": true_answer_images, "correct": True},
                        *[{"text": fa["text"], "images": fa["images"], "correct": False} for fa in false_answers]
                    ]
                    options_data = [opt for opt in options_data if opt["text"] or opt["images"]]
                    if len(options_data) < 4:
                        continue
                    correct_index = next((i for i, opt in enumerate(options_data) if opt["correct"]), 0)
                    if self.settings["shuffle_answers"]:
                        random.shuffle(options_data)
                        correct_index = next((i for i, opt in enumerate(options_data) if opt["correct"]), 0)
                    questions.append({
                        "question": question_text,
                        "question_images": question_images,
                        "options": [{"text": opt["text"], "images": opt["images"]} for opt in options_data],
                        "correct": correct_index
                    })
            return questions
        except Exception as e:
            messagebox.showerror("Xato", f"Test faylini yuklashda xato: {str(e)}")
            return []

    def extract_images_from_element(self, element, doc):
        images = []
        try:
            for rel in element.xpath('.//a:blip'):
                rId = rel.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                if rId in doc.part.related_parts:
                    image_part = doc.part.related_parts[rId]
                    image_data = image_part.blob
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp:
                        tmp.write(image_data)
                        images.append(tmp.name)
                        self.temp_files.append(tmp.name)
        except Exception as e:
            print(f"Rasmlarni olishda xato: {e}")
        return images

    def start_new_test(self):
        try:
            # Agar oldingi testdan tahlil ma'lumotlari bo'lsa, yuborish
            if hasattr(self, 'review_start_time') and self.review_start_time:
                self.send_review_status(force_send=True)
            
            if not self.all_questions:
                messagebox.showerror("Xato", "Test savollari yuklanmagan!")
                return
                
            # Test boshlashdan oldin tozalash
            self.cleanup_previous_test()
                
            all_indices = set(range(len(self.all_questions)))
            available_indices = list(all_indices - self.used_question_indices)
            
            if len(available_indices) < self.settings["questions_per_test"]:
                available_indices = list(all_indices)
                self.used_question_indices = set()
                
            selected_indices = random.sample(
                available_indices, 
                min(self.settings["questions_per_test"], len(available_indices))
            )
            
            self.current_test_set = [self.all_questions[i] for i in selected_indices]
            self.used_question_indices.update(selected_indices)
            
            # Test oynasini ko'rsatish
            self.show_test_screen()
            self.start_timer()
            
        except Exception as e:
            messagebox.showerror("Xato", f"Yangi test boshlashda xatolik: {str(e)}")
            return

    def cleanup_previous_test(self):
        """Oldingi testdan qolgan ma'lumotlarni tozalash"""
        try:
            # Timer va test holatini tozalash
            self.timer_running = False
            self.timer_seconds = self.settings["test_time"] * 60
            
            # Test ma'lumotlarini tozalash
            self.user_answers = {i: None for i in range(self.settings["questions_per_test"])}
            self.current_question = 0
            self.admin_override = False
            self.results_shown = False
            self.current_zoom = 1.0
            
            # Rasmlar va vaqtinchalik fayllarni tozalash
            if hasattr(self, 'images'):
                self.images = []
            
            # Oynalarni tozalash
            if hasattr(self, 'zoom_window') and self.zoom_window:
                try:
                    self.zoom_window.destroy()
                    self.zoom_window = None
                except:
                    pass
                    
        except Exception as e:
            print(f"Tozalashda xatolik: {str(e)}")
            # Xato bo'lsa ham davom etamiz

    def restart_current_test(self):
        self.user_answers = {i: None for i in range(self.settings["questions_per_test"])}
        self.current_question = 0
        self.timer_seconds = self.settings["test_time"] * 60
        self.timer_running = True
        self.admin_override = False
        self.results_shown = False
        self.current_zoom = 1.0
        self.show_test_screen()
        self.start_timer()

    def show_test_screen(self):
        try:
            if not self.current_test_set:
                messagebox.showerror("Xato", "Test savollari topilmadi!")
                return
                
            self.current_screen = "test"
            self.clear_screen()
            
            # Timer ni qayta o'rnatish
            if not hasattr(self, 'timer_seconds') or self.timer_seconds <= 0:
                self.timer_seconds = self.settings["test_time"] * 60
            self.timer_running = True
            
            header_frame = ttk.Frame(self.root, padding=10)
            header_frame.pack(fill='x', pady=(10, 0))
            # Foydalanuvchi ismi va familiyasi
            ttk.Label(
                header_frame,
                text=f"Foydalanuvchi: {self.username}",
                font=('Segoe UI', 13, 'bold'),
                foreground=self.colors['primary']
            ).pack(side='right', padx=10)
            timer_frame = ttk.Frame(header_frame)
            timer_frame.pack(side='left', padx=10)
            ttk.Label(
                timer_frame, 
                text="Qolgan vaqt:", 
                font=('Segoe UI', 12, 'bold')
            ).pack(side='left')
            
            # Timer labelni to'g'rilash
            mins, secs = divmod(self.timer_seconds, 60)
            self.timer_label = ttk.Label(
                timer_frame, 
                text=f"{mins:02}:{secs:02}", 
                font=('Segoe UI', 12, 'bold'),
                foreground=self.colors['primary']
            )
            self.timer_label.pack(side='left', padx=5)
            
            # Progress bar yangilash
            answered_count = sum(1 for ans in self.user_answers.values() if ans is not None)
            self.progress = ttk.Progressbar(
                header_frame, 
                length=400, 
                maximum=self.settings["questions_per_test"], 
                mode='determinate',
                style='Progress.Horizontal.TProgressbar',
                value=answered_count
            )
            self.progress.pack(side='right', padx=10)
            
            self.question_number_label = ttk.Label(
                header_frame,
                text=f"Savol: {self.current_question+1}/{self.settings['questions_per_test']}",
                font=('Segoe UI', 13, 'bold'),
                foreground=self.colors['secondary']
            )
            self.question_number_label.pack(side='left', padx=20)
            
            main_card = ttk.Frame(
                self.root, 
                style='Card.TFrame',
                padding=20
            )
            main_card.pack(pady=10, padx=20, fill='both', expand=True)
            self.question_frame = ttk.Frame(main_card)
            self.question_frame.pack(fill='both', expand=True)
            self.option_vars = tk.IntVar(value=-1)
            self.option_buttons = []
            nav_control_frame = ttk.Frame(self.root, padding=10)
            nav_control_frame.pack(fill='x')
            ttk.Button(
                nav_control_frame, 
                text="◀ Oldingi", 
                style='Secondary.TButton',
                command=self.previous_question
            ).pack(side='left', padx=5)
            self.next_btn = ttk.Button(
                nav_control_frame, 
                text="Keyingi ▶", 
                style='Secondary.TButton',
                command=self.next_question
            )
            self.next_btn.pack(side='left', padx=5)
            ttk.Button(
                nav_control_frame,
                text="Yordam",
                style='Accent.TButton',
                command=self.show_help
            ).pack(side='right', padx=5)
            ttk.Button(
                nav_control_frame,
                text="Admin Panel",
                style='Accent.TButton',
                command=self.show_admin_panel
            ).pack(side='right', padx=5)
            nav_card = ttk.Frame(
                self.root, 
                style='Card.TFrame',
                padding=10
            )
            nav_card.pack(pady=10, padx=20, fill='x')
            self.nav_buttons = []
            questions_per_row = 15
            for i in range(self.settings["questions_per_test"]):
                btn = ttk.Button(
                    nav_card, 
                    text=str(i+1), 
                    width=3,
                    command=lambda i=i: self.go_to_question(i)
                )
                btn.grid(
                    row=i//questions_per_row, 
                    column=i%questions_per_row, 
                    padx=2, 
                    pady=2
                )
                self.nav_buttons.append(btn)
            control_frame = ttk.Frame(self.root, padding=10)
            control_frame.pack(fill='x')
            self.new_test_button = ttk.Button(
                control_frame, 
                text="Yangi Test", 
                style='Secondary.TButton',
                command=self.admin_new_test, 
                state='normal'
            )
            self.new_test_button.grid(row=0, column=0, padx=5)
            self.restart_button = ttk.Button(
                control_frame, 
                text="Qayta Boshlash", 
                style='Secondary.TButton',
                command=self.admin_restart_test, 
                state='normal'
            )
            self.restart_button.grid(row=0, column=1, padx=5)
            ttk.Button(
                control_frame, 
                text="Yakunlash", 
                style='Success.TButton',
                command=self.normal_finish
            ).grid(row=0, column=2, padx=5)
            ttk.Button(
                control_frame, 
                text="Admin Yakunlash", 
                style='Warning.TButton',
                command=self.admin_finish
            ).grid(row=0, column=3, padx=5)
            self.progress_label = ttk.Label(
                control_frame, 
                text=f"Progress: 1/{self.settings['questions_per_test']}", 
                font=('Segoe UI', 11)
            )
            self.progress_label.grid(row=0, column=4, padx=10, sticky='e')
            self.show_question(self.current_question)
            self.update_nav_buttons()
            # Klaviatura bog'lanishlarini qo'shish
            self.root.bind('<Left>', lambda event: self.previous_question())
            self.root.bind('<Right>', lambda event: self.next_question())
            self.root.bind('a', lambda event: self.select_option_by_key(0))
            self.root.bind('b', lambda event: self.select_option_by_key(1))
            self.root.bind('c', lambda event: self.select_option_by_key(2))
            self.root.bind('d', lambda event: self.select_option_by_key(3))
            
        except Exception as e:
            messagebox.showerror("Xato", f"Test oynasini yaratishda xato: {str(e)}")

    def show_help(self):
        help_win = Toplevel(self.root)
        help_win.title("Yordam")
        help_win.geometry("500x450")  # Balandlikni oshirdik
        q = self.current_test_set[self.current_question]
        source_file = q.get('source_file', 'Nomaʼlum')
        source_index = q.get('source_index', 'Nomaʼlum')
        help_text = (
            "Test dasturidan foydalanish bo'yicha qisqacha qo'llanma:\n"
            "- Savollar va variantlar ekranda ko'rsatiladi.\n"
            "- Variantni tanlash uchun ustiga bosing yoki klaviaturada A/B/C/D tugmasini bosing.\n"
            "- Keyingi yoki oldingi savolga o'tish uchun tugmalardan yoki ←/→ tugmalaridan foydalaning.\n"
            "- Har bir savol raqamiga bosib, istalgan savolga o'tishingiz mumkin.\n"
            "- Test yakunida natijani ko'rasiz va har bir savolni qayta ko'rib chiqishingiz mumkin.\n"
            "- Qorong'i/Yorug' rejimni sozlamalardan o'zgartiring.\n"
            f"\n\n**Hozirgi savol:**\n- Baza fayli: {source_file}\n- Fayldagi savol raqami: {source_index}"
        )
        label = ttk.Label(help_win, text=help_text, font=('Segoe UI', 12), wraplength=480, justify='left')
        label.pack(padx=20, pady=20)
        
        # Tahrirlash tugmasini qo'shish
        def edit_current_question():
            if not messagebox.askyesno("Tasdiqlash", 
                "Bu savolni tahrirlash uchun admin huquqi kerak. Davom etasizmi?"):
                return
                
            password = simpledialog.askstring(
                "Admin tasdiqlash",
                "Admin parolini kiriting:",
                show='*',
                parent=help_win
            )
            
            if password != self.admin_password:
                messagebox.showerror("Xato", "Noto'g'ri admin paroli!")
                return
                
            # Admin panelni ochish va joriy savolni yuklash
            help_win.destroy()
            admin_win = Toplevel(self.root)
            admin_win.title("Admin Panel - Savolni tahrirlash")
            admin_win.geometry("800x700")
            
            # Joriy savol ma'lumotlarini olish
            q = self.current_test_set[self.current_question]
            source_file = q.get('source_file', '')
            source_index = q.get('source_index', '')
            
            # Admin panel elementlarini yaratish
            files = glob.glob(os.path.join(BASE_DIR, "baza", "*.docx"))
            file_names = [os.path.basename(f) for f in files]
            file_var = tk.StringVar(value=source_file if source_file in file_names else file_names[0])
            file_combo = ttk.Combobox(admin_win, values=file_names, textvariable=file_var, state='readonly', font=('Segoe UI', 12))
            file_combo.pack(pady=5)
            
            # Savol raqami
            ttk.Label(admin_win, text="Savol raqami:", font=('Segoe UI', 11)).pack(pady=(10, 0))
            index_entry = ttk.Entry(admin_win, font=('Segoe UI', 12))
            index_entry.insert(0, str(source_index))
            index_entry.pack(pady=5)
            
            # Savol matni
            ttk.Label(admin_win, text="Savol matni:", font=('Segoe UI', 11)).pack(pady=(10, 0))
            question_entry = tk.Text(admin_win, height=3, font=('Segoe UI', 12), wrap='word')
            question_entry.insert('1.0', q.get('question', ''))
            question_entry.pack(pady=5, fill='x', padx=10)
            
            # Variantlar
            variant_entries = []
            for i, opt in enumerate(q.get('options', [])):
                frame = ttk.Frame(admin_win)
                frame.pack(fill='x', padx=10, pady=2)
                ttk.Label(frame, text=f"Variant {chr(65+i)}:", font=('Segoe UI', 11)).pack(side='left')
                entry = ttk.Entry(frame, font=('Segoe UI', 12), width=60)
                entry.insert(0, opt.get('text', ''))
                entry.pack(side='left', padx=5)
                variant_entries.append(entry)
            
            # To'g'ri javob
            correct_var = tk.IntVar(value=q.get('correct', 0))
            correct_frame = ttk.Frame(admin_win)
            correct_frame.pack(pady=5)
            ttk.Label(correct_frame, text="To'g'ri javob:", font=('Segoe UI', 11)).pack(side='left')
            for i in range(4):
                ttk.Radiobutton(correct_frame, text=chr(65+i), variable=correct_var, value=i).pack(side='left', padx=5)
            
            # Tugmalar
            btn_frame = ttk.Frame(admin_win)
            btn_frame.pack(pady=10)
            ttk.Button(btn_frame, text="Saqlash", style='Primary.TButton',
                      command=lambda: self.admin_edit_question(file_var.get(), index_entry, question_entry, variant_entries, correct_var, admin_win)).pack(side='left', padx=5)
            ttk.Button(btn_frame, text="Yopish", style='Danger.TButton',
                      command=admin_win.destroy).pack(side='left', padx=5)
        
        btn_frame = ttk.Frame(help_win)
        btn_frame.pack(pady=10)
        
        ttk.Button(btn_frame, text="Shu savolni tahrirlash", style='Warning.TButton',
                  command=edit_current_question).pack(side='left', padx=5)
        ttk.Button(btn_frame, text="Yopish", style='Primary.TButton',
                  command=help_win.destroy).pack(side='left', padx=5)

    def select_option_by_key(self, idx):
        """Klaviatura orqali variant tanlash"""
        try:
            if not hasattr(self, 'current_test_set') or not self.current_test_set:
                return
            if self.current_question >= len(self.current_test_set):
                return
            current_question = self.current_test_set[self.current_question]
            if not current_question or 'options' not in current_question:
                return
            if idx < len(current_question['options']):
                self.option_vars.set(idx)
                self.select_answer()
        except Exception as e:
            print(f"Klaviatura bilan tanlashda xatolik: {e}")

    def show_question(self, index):
        self.current_question = index
        for widget in self.question_frame.winfo_children():
            widget.destroy()
        question_data = self.current_test_set[index]
        question_text = f"{index+1}. {question_data['question']}"
        container = ttk.Frame(self.question_frame)
        container.pack(anchor='center', fill='both', expand=True)
        ttk.Label(
            container, 
            text=question_text, 
            wraplength=900, 
            font=('Segoe UI', 14), 
            justify='left',
            background='white'
        ).pack(anchor='w', pady=(0, 10))
        if question_data.get("question_images"):
            img_frame = ttk.Frame(container)
            img_frame.pack(pady=10)
            for img_path in question_data['question_images']:
                try:
                    img = Image.open(img_path)
                    max_width, max_height = 300, 200
                    img.thumbnail((max_width, max_height), Image.Resampling.LANCZOS)
                    if self.settings['dark_mode']:
                        border_color = self.colors['border']
                    else:
                        border_color = (222, 226, 230)
                    bordered_img = Image.new('RGB', 
                                           (img.width + 10, img.height + 10), 
                                           border_color)
                    bordered_img.paste(img, (5, 5))
                    photo = ImageTk.PhotoImage(bordered_img)
                    self.images.append(photo)
                    img_label = ttk.Label(img_frame, image=photo)
                    img_label.pack(side='left', padx=5)
                    zoom_btn = ttk.Button(
                        img_frame, 
                        text="Zoom", 
                        style='Accent.TButton',
                        command=lambda path=img_path: self.show_zoom_window(path)
                    )
                    zoom_btn.pack(side='left', padx=5)
                except Exception as e:
                    print(f"Savol rasmini yuklashda xato: {e}")
        self.option_vars.set(self.user_answers[index] if self.user_answers[index] is not None else -1)
        for i, opt in enumerate(question_data['options']):
            option_frame = ttk.Frame(container)
            option_frame.pack(anchor='w', fill='x', pady=5)
            label = chr(65 + i)  # 'A', 'B', 'C', 'D'
            btn = tk.Radiobutton(
                option_frame, 
                text=f"{label}) {opt['text']}", 
                variable=self.option_vars,
                value=i, 
                font=('Segoe UI', 12), 
                anchor='w', 
                justify='left',
                wraplength=800,
                bg='white',
                activebackground=self.colors['light'],
                selectcolor=self.colors['accent'],
                relief='flat',
                padx=10,
                command=self.select_answer
            )
            btn.pack(side='left', anchor='w', fill='x', expand=True)
            if opt.get('images'):
                opt_img_frame = ttk.Frame(option_frame)
                opt_img_frame.pack(side='left')
                for img_path in opt['images']:
                    try:
                        img = Image.open(img_path)
                        max_width, max_height = 100, 70
                        img.thumbnail((max_width, max_height), Image.Resampling.LANCZOS)
                        bordered_img = Image.new('RGB', 
                                               (img.width + 4, img.height + 4), 
                                               self.colors['border'])
                        bordered_img.paste(img, (2, 2))
                        photo = ImageTk.PhotoImage(bordered_img)
                        self.images.append(photo)
                        img_label = ttk.Label(opt_img_frame, image=photo)
                        img_label.pack(side='left', padx=2)
                    except Exception as e:
                        print(f"Variant rasmini yuklashda xato: {e}")
        self.progress_label.config(
            text=f"Progress: {index+1}/{self.settings['questions_per_test']}"
        )
        self.progress['value'] = index + 1
        self.question_number_label.config(
            text=f"Savol: {index+1}/{self.settings['questions_per_test']}"
        )
        for i, btn in enumerate(self.nav_buttons):
            if i == index:
                btn.configure(style='Accent.TButton')
            elif self.user_answers.get(i) is not None:
                btn.configure(style='Success.TButton')
            else:
                btn.configure(style='TButton')

    def select_answer(self):
        selected = self.option_vars.get()
        if selected == -1:
            return
        self.user_answers[self.current_question] = selected
        self.update_nav_buttons()

    def update_nav_buttons(self):
        for i, widget in enumerate(self.nav_buttons):
            if self.user_answers.get(i) is not None:
                widget.configure(style='Accent.TButton')
            else:
                widget.configure(style='TButton')

    def go_to_question(self, index):
        self.select_answer()
        self.show_question(index)

    def previous_question(self):
        if self.current_question > 0:
            self.go_to_question(self.current_question - 1)

    def next_question(self):
        if self.current_question < self.settings["questions_per_test"] - 1:
            self.go_to_question(self.current_question + 1)

    def start_timer(self):
        if not self.timer_running:
            return
            
        if self.timer_seconds <= 0:
            self.normal_finish()
            return
            
        try:
            mins, secs = divmod(self.timer_seconds, 60)
            if hasattr(self, 'timer_label') and self.timer_label.winfo_exists():
                self.timer_label.config(text=f"{mins:02}:{secs:02}")
                if self.timer_seconds <= 60:
                    self.timer_label.config(foreground=self.colors['danger'])
                else:
                    self.timer_label.config(foreground=self.colors['primary'])
            
            self.timer_seconds -= 1
            if self.timer_running:
                self.root.after(1000, self.start_timer)
                
            # Progress bar va label yangilash
            if hasattr(self, 'progress') and self.progress.winfo_exists():
                answered_count = sum(1 for ans in self.user_answers.values() if ans is not None)
                self.progress['value'] = answered_count
                if hasattr(self, 'progress_label') and self.progress_label.winfo_exists():
                    self.progress_label.config(
                        text=f"Javob berilgan: {answered_count}/{self.settings['questions_per_test']}"
                    )
                    
        except Exception as e:
            print(f"Timer xatosi: {str(e)}")
            self.timer_running = False

    def normal_finish(self):
        answered = sum(1 for ans in self.user_answers.values() if ans is not None)
        if answered < self.settings["questions_per_test"]:
            messagebox.showwarning(
                "Diqqat",
                f"Testni yakunlash uchun barcha savollarga javob bering!\nJavob berilgan: {answered}/{self.settings['questions_per_test']}"
            )
            return
        self.finish_test()

    def admin_finish(self):
        password = simpledialog.askstring(
            "Parol", 
            "Admin parolini kiriting:", 
            show='*'
        )
        if password == self.admin_password:
            self.admin_override = True
            self.finish_test()
        else:
            messagebox.showerror("Xato", "Noto'g'ri parol!")

    def finish_test(self):
        try:
            # Timer to'xtatish
            self.timer_running = False
            self.timer_seconds = 0
            
            self.current_screen = "results"
            self.results_shown = True
            
            # Klaviatura bog'lanishlarini o'chirish
            self.unbind_keyboard_shortcuts()
            
            if hasattr(self, 'new_test_button') and self.new_test_button.winfo_exists():
                self.new_test_button.config(state='normal')
            if hasattr(self, 'restart_button') and self.restart_button.winfo_exists():
                self.restart_button.config(state='normal')
                
            answered = sum(1 for ans in self.user_answers.values() if ans is not None)
            score = 0
            wrong_answers = []
            for i, q in enumerate(self.current_test_set):
                if self.user_answers.get(i) == q['correct']:
                    score += 1
                else:
                    wrong_answers.append(i)

            percentage = round(score / self.settings["questions_per_test"] * 100, 2)
            
            # Xatolarni tahlil qilish vaqtini boshlash
            if not hasattr(self, 'review_start_time') or not self.review_start_time:
                self.review_start_time = datetime.now()
                self.reviewed_questions = set()
                self.wrong_answers = wrong_answers
                self.initial_result_sent = False
            
            # Ekranni tozalash va yangi ma'lumotlarni ko'rsatish
            for widget in self.root.winfo_children():
                widget.destroy()
                
            main_frame = ttk.Frame(self.root, padding=20)
            main_frame.pack(expand=True, fill='both')
            
            ttk.Label(
                main_frame, 
                text=f"{self.username}, natijangiz:", 
                font=('Segoe UI', 24, 'bold'),
                foreground=self.colors['primary']
            ).pack(pady=(0, 20))
            
            results_card = ttk.Frame(
                main_frame, 
                style='Card.TFrame',
                padding=30
            )
            results_card.pack(pady=10, fill='x')
            
            metrics = [
                ("To'g'ri javoblar soni:", f"{score}/{self.settings['questions_per_test']}"),
                ("Foiz:", f"{percentage}%"),
                ("Javob berilgan savollar:", f"{answered}/{self.settings['questions_per_test']}")
            ]
            
            for metric, value in metrics:
                metric_frame = ttk.Frame(results_card)
                metric_frame.pack(fill='x', pady=5)
                ttk.Label(
                    metric_frame, 
                    text=metric, 
                    font=('Segoe UI', 14),
                    width=25,
                    anchor='w'
                ).pack(side='left')
                ttk.Label(
                    metric_frame, 
                    text=value, 
                    font=('Segoe UI', 14, 'bold'),
                    foreground=self.colors['primary']
                ).pack(side='left')
                
            # Filtrlash tugmalari
            filter_frame = ttk.Frame(main_frame)
            filter_frame.pack(pady=10)
            ttk.Button(
                filter_frame,
                text="Barchasi",
                style='Primary.TButton',
                command=lambda: self.show_review_filtered(None)
            ).pack(side='left', padx=5)
            ttk.Button(
                filter_frame,
                text="Faqat noto'g'ri javoblar",
                style='Danger.TButton',
                command=lambda: self.show_review_filtered('wrong')
            ).pack(side='left', padx=5)
            ttk.Button(
                filter_frame,
                text="Javob bermaganlar",
                style='Warning.TButton',
                command=lambda: self.show_review_filtered('empty')
            ).pack(side='left', padx=5)
            
            # review_card faqat bitta frame bo'ladi
            self.review_card = ttk.Frame(
                main_frame, 
                style='Card.TFrame',
                padding=20
            )
            self.review_card.pack(pady=20, fill='x')
            
            # Darhol ko'rsatish
            self.root.update_idletasks()
            self.show_review_filtered(None)
            
            action_frame = ttk.Frame(main_frame)
            action_frame.pack(pady=20)
            ttk.Button(
                action_frame, 
                text="Bosh sahifa", 
                style='Primary.TButton',
                command=self.return_to_main_menu
            ).pack(side='left', padx=10, ipadx=20)
            ttk.Button(
                action_frame, 
                text="Dasturdan Chiqish", 
                style='Danger.TButton',
                command=self.on_close
            ).pack(side='left', padx=10, ipadx=20)
            
            if hasattr(self, 'admin_override'):
                del self.admin_override
                
            # Test natijasini saqlash va yuborish (faqat birinchi marta)
            if not hasattr(self, 'initial_result_sent') or not self.initial_result_sent:
                self.root.after(100, lambda: self.save_test_result(
                    student_name=self.username,
                    score=score,
                    total_questions=self.settings["questions_per_test"]
                ))
                
                # Telegram botga screenshot yuborish (faqat birinchi marta)
                try:
                    chat_id = None
                    if os.path.exists('telegram_chat_id.txt'):
                        with open('telegram_chat_id.txt', 'r') as f:
                            chat_id = f.read().strip()
                    global global_admin_chat_id
                    if not chat_id and global_admin_chat_id:
                        chat_id = global_admin_chat_id
                    if not chat_id:
                        chat_id = get_latest_chat_id(TELEGRAM_BOT_TOKEN)
                        if chat_id:
                            with open('telegram_chat_id.txt', 'w') as f:
                                f.write(chat_id)
                            global_admin_chat_id = chat_id
                    if chat_id:
                        result_text = (
                            f"Foydalanuvchi: {self.username}\n"
                            f"To'g'ri javoblar: {score}/{self.settings['questions_per_test']}\n"
                            f"Foiz: {percentage}%\n"
                        )
                        # Screenshot olish va yuborish
                        was_fullscreen = self.root.attributes('-fullscreen')
                        self.root.attributes('-fullscreen', True)
                        self.root.lift()
                        self.root.attributes('-topmost', True)
                        self.root.update_idletasks()
                        self.root.update()
                        time.sleep(0.3)
                        screenshot = ImageGrab.grab()
                        self.root.attributes('-topmost', False)
                        if was_fullscreen and (was_fullscreen == 1 or was_fullscreen == True or was_fullscreen == '1'):
                            self.root.attributes('-fullscreen', True)
                        else:
                            self.root.attributes('-fullscreen', False)
                        self.root.update()
                        temp_path = os.path.join(tempfile.gettempdir(), f"test_result_{int(time.time())}.png")
                        screenshot.save(temp_path)
                        send_result_screenshot_to_telegram(TELEGRAM_BOT_TOKEN, chat_id, temp_path, caption=result_text)
                        try:
                            os.remove(temp_path)
                        except Exception:
                            pass
                except Exception as e:
                    print(f"Telegramga natija yuborishda xato: {e}")
                
                self.initial_result_sent = True
                
            # Oynani yangilash
            self.root.update_idletasks()
            self.root.update()
            
        except Exception as e:
            print(f"Test yakunlashda xato: {str(e)}")
            traceback.print_exc()
            messagebox.showerror("Xato", f"Test yakunlashda xatolik yuz berdi: {str(e)}")
            # Xatolik yuz berganda ham asosiy oynani ko'rsatishga harakat qilish
            try:
                self.create_start_screen()
            except:
                pass

    def show_review_filtered(self, filter_type):
        # Faqat self.review_card ichini tozalash
        for widget in self.review_card.winfo_children():
            widget.destroy()
        questions_per_row = 10
        for i, q in enumerate(self.current_test_set):
            user_ans = self.user_answers.get(i)
            if filter_type == 'wrong' and (user_ans is None or user_ans == q['correct']):
                continue
            if filter_type == 'empty' and user_ans is not None:
                continue
            if user_ans == q['correct']:
                style = 'Success.TButton'
            elif user_ans is not None:
                style = 'Danger.TButton'
            else:
                style = 'Warning.TButton'
            btn = ttk.Button(
                self.review_card, 
                text=str(i+1), 
                style=style,
                command=lambda i=i: self.review_question(i)
            )
            btn.grid(
                row=i//questions_per_row, 
                column=i%questions_per_row, 
                padx=2, 
                pady=2
            )

    def review_question(self, index):
        # Xatolarni tahlil qilish kuzatuvini qo'shish
        self.reviewed_questions.add(index)
        self.clear_screen()
        question_data = self.current_test_set[index]
        main_frame = ttk.Frame(self.root, padding=20)
        main_frame.pack(expand=True, fill='both')
        ttk.Label(
            main_frame, 
            text=f"{index+1}. {question_data['question']}",
            wraplength=900, 
            font=('Segoe UI', 16), 
            anchor='w', 
            justify='left',
            background='white'
        ).pack(anchor='w', pady=(0, 10))
        if question_data.get("question_images"):
            img_frame = ttk.Frame(main_frame)
            img_frame.pack(pady=10)
            for img_path in question_data['question_images']:
                try:
                    img = Image.open(img_path)
                    max_width, max_height = 300, 200
                    img.thumbnail((max_width, max_height), Image.Resampling.LANCZOS)
                    bordered_img = Image.new('RGB', 
                                           (img.width + 10, img.height + 10), 
                                           self.colors['border'])
                    bordered_img.paste(img, (5, 5))
                    photo = ImageTk.PhotoImage(bordered_img)
                    self.images.append(photo)
                    img_label = ttk.Label(img_frame, image=photo)
                    img_label.pack(side='left', padx=5)
                    zoom_btn = ttk.Button(
                        img_frame, 
                        text="Zoom", 
                        style='Accent.TButton',
                        command=lambda path=img_path: self.show_zoom_window(path)
                    )
                    zoom_btn.pack(side='left', padx=5)
                except Exception as e:
                    print(f"Savol rasmini yuklashda xato: {e}")
        for i, opt in enumerate(question_data['options']):
            if i == question_data['correct']:
                color = self.colors['success']
                weight = 'bold'
            elif i == self.user_answers.get(index):
                color = self.colors['danger']
                weight = 'bold'
            else:
                color = self.colors['text']
                weight = 'normal'
            option_frame = ttk.Frame(main_frame)
            option_frame.pack(anchor='w', fill='x', pady=5)
            label = chr(65 + i)
            ttk.Label(
                option_frame, 
                text=f"{label}) {opt['text']}", 
                foreground=color, 
                font=('Segoe UI', 12, weight), 
                anchor='w', 
                justify='left',
                wraplength=800,
                background='white'
            ).pack(side='left', anchor='w')
            if opt.get('images'):
                opt_img_frame = ttk.Frame(option_frame)
                opt_img_frame.pack(side='left')
                for img_path in opt['images']:
                    try:
                        img = Image.open(img_path)
                        max_width, max_height = 100, 70
                        img.thumbnail((max_width, max_height), Image.Resampling.LANCZOS)
                        border_width = 3 if color != self.colors['text'] else 1
                        bordered_img = Image.new('RGB', 
                                               (img.width + border_width*2, img.height + border_width*2), 
                                               color if color != self.colors['text'] else self.colors['border'])
                        bordered_img.paste(img, (border_width, border_width))
                        photo = ImageTk.PhotoImage(bordered_img)
                        self.images.append(photo)
                        img_label = ttk.Label(opt_img_frame, image=photo)
                        img_label.pack(side='left', padx=2)
                    except Exception as e:
                        print(f"Variant rasmini yuklashda xato: {e}")
        nav_frame = ttk.Frame(main_frame)
        nav_frame.pack(pady=20)
        questions_per_row = 10
        for i, q in enumerate(self.current_test_set):
            user_ans = self.user_answers.get(i)
            if user_ans == q['correct']:
                style = 'Success.TButton'
            elif user_ans is not None:
                style = 'Danger.TButton'
            else:
                style = 'Warning.TButton'
            btn = ttk.Button(
                nav_frame, 
                text=str(i+1), 
                style=style,
                command=lambda i=i: self.review_question(i)
            )
            btn.grid(
                row=i//questions_per_row, 
                column=i%questions_per_row, 
                padx=2, 
                pady=2
            )
        ttk.Button(
            main_frame, 
            text="Natijalarga qaytish", 
            style='Primary.TButton',
            command=self.show_results
        ).pack(pady=10)

    def show_results(self):
        # Natijalarga qaytishda xatolar tahlilini yubormaslik uchun vaqtincha o'zgaruvchini saqlash
        temp_review_time = self.review_start_time
        temp_reviewed = self.reviewed_questions.copy() if hasattr(self, 'reviewed_questions') else set()
        temp_wrong = self.wrong_answers if hasattr(self, 'wrong_answers') else []
        temp_initial_sent = self.initial_result_sent if hasattr(self, 'initial_result_sent') else False
        
        # Natijalar oynasini ko'rsatish
        self.finish_test()
        
        # Tahlil ma'lumotlarini qayta tiklash
        self.review_start_time = temp_review_time
        self.reviewed_questions = temp_reviewed
        self.wrong_answers = temp_wrong
        self.initial_result_sent = temp_initial_sent

    def clear_screen(self):
        for widget in self.root.winfo_children():
            widget.destroy()
        self.images = []
        if hasattr(self, 'zoom_window') and self.zoom_window:
            self.zoom_window.destroy()
            self.zoom_window = None
        # Klaviatura bog'lanishlarini o'chirish
        self.unbind_keyboard_shortcuts()

    def unbind_keyboard_shortcuts(self):
        """Klaviatura bog'lanishlarini o'chirish"""
        try:
            self.root.unbind('<Left>')
            self.root.unbind('<Right>')
            self.root.unbind('a')
            self.root.unbind('b')
            self.root.unbind('c')
            self.root.unbind('d')
            self.root.unbind('<space>')
        except Exception as e:
            print(f"Klaviatura bog'lanishlarini o'chirishda xatolik: {e}")

    def show_zoom_window(self, image_path):
        if self.zoom_window and self.zoom_window.winfo_exists():
            self.zoom_window.destroy()
        self.zoom_window = Toplevel(self.root)
        self.zoom_window.title("Zoom Rasm")
        self.zoom_window.geometry("800x600")
        try:
            self.original_img = Image.open(image_path)
            self.current_zoom = 1.0
            win_width, win_height = 800, 600
            self.canvas = tk.Canvas(
                self.zoom_window, 
                width=win_width, 
                height=win_height,
                scrollregion=(0, 0, win_width, win_height),
                bg=self.colors['light']
            )
            h_scroll = ttk.Scrollbar(
                self.zoom_window, 
                orient='horizontal', 
                command=self.canvas.xview
            )
            v_scroll = ttk.Scrollbar(
                self.zoom_window, 
                orient='vertical', 
                command=self.canvas.yview
            )
            self.canvas.configure(
                xscrollcommand=h_scroll.set, 
                yscrollcommand=v_scroll.set
            )
            h_scroll.pack(side='bottom', fill='x')
            v_scroll.pack(side='right', fill='y')
            self.canvas.pack(fill='both', expand=True)
            self.img_frame = ttk.Frame(self.canvas)
            self.canvas.create_window((0, 0), window=self.img_frame, anchor='nw')
            self.update_zoomed_image()
            btn_frame = ttk.Frame(self.zoom_window)
            btn_frame.pack(pady=5)
            ttk.Button(
                btn_frame, 
                text="Kattalashtirish (+)",
                style='Accent.TButton',
                command=lambda: self.zoom_image(1.2)
            ).pack(side='left', padx=5)
            ttk.Button(
                btn_frame, 
                text="Kichiklashtirish (-)",
                style='Accent.TButton',
                command=lambda: self.zoom_image(0.8)
            ).pack(side='left', padx=5)
            self.zoom_window.bind("<MouseWheel>", self.on_mouse_wheel)
        except Exception as e:
            messagebox.showerror("Xato", f"Rasmni ochishda xato: {e}")
            self.zoom_window.destroy()

    def update_zoomed_image(self):
        if not hasattr(self, 'original_img'):
            return
        width = int(self.original_img.width * self.current_zoom)
        height = int(self.original_img.height * self.current_zoom)
        img = self.original_img.copy()
        img = img.resize((width, height), Image.Resampling.LANCZOS)
        bordered_img = Image.new('RGB', 
                               (img.width + 10, img.height + 10), 
                               self.colors['border'])
        bordered_img.paste(img, (5, 5))
        self.zoomed_image = ImageTk.PhotoImage(bordered_img)
        for widget in self.img_frame.winfo_children():
            widget.destroy()
        self.img_label = ttk.Label(self.img_frame, image=self.zoomed_image)
        self.img_label.pack()
        self.canvas.configure(scrollregion=self.canvas.bbox('all'))

    def zoom_image(self, factor):
        self.current_zoom *= factor
        self.current_zoom = max(0.1, min(5.0, self.current_zoom))
        self.update_zoomed_image()

    def on_mouse_wheel(self, event):
        if event.delta > 0:
            self.zoom_image(1.1)
        else:
            self.zoom_image(0.9)

    def on_close(self):
        if messagebox.askokcancel("Chiqish", "Dasturdan chiqishni istaysizmi?"):
            # Agar test yakunlangan bo'lsa va xatolar ko'rib chiqilmagan bo'lsa
            if hasattr(self, 'review_start_time') and self.review_start_time:
                self.send_review_status(force_send=True)
            self.cleanup()
            self.root.destroy()
            sys.exit()

    def admin_new_test(self):
        password = simpledialog.askstring(
            "Admin paroli",
            "Yangi test boshlash uchun admin parolini kiriting:",
            show='*',
            parent=self.root
        )
        if password == self.admin_password:
            self.start_new_test()
        else:
            messagebox.showerror("Xato", "Noto'g'ri admin paroli! Yangi test boshlanmadi.")

    def admin_restart_test(self):
        password = simpledialog.askstring(
            "Admin paroli",
            "Testni qayta boshlash uchun admin parolini kiriting:",
            show='*',
            parent=self.root
        )
        if password == self.admin_password:
            self.restart_current_test()
        else:
            messagebox.showerror("Xato", "Noto'g'ri admin paroli! Qayta boshlanmadi.")

    def show_admin_panel(self):
        password = simpledialog.askstring(
            "Admin paroli",
            "Admin panelga kirish uchun parolni kiriting:",
            show='*',
            parent=self.root
        )
        if password != self.admin_password:
            messagebox.showerror("Xato", "Noto'g'ri admin paroli!")
            return
        admin_win = Toplevel(self.root)
        admin_win.title("Admin Panel - Testlarni boshqarish")
        admin_win.geometry("800x700")
        admin_win.grab_set()
        # --- TUGMALAR ---
        btns_frame = ttk.Frame(admin_win)
        btns_frame.pack(fill='x', padx=10, pady=10)
        ttk.Button(btns_frame, text="Sozlamalar", style='Primary.TButton', 
                   command=lambda: self.show_admin_settings(admin_win)).pack(side='left', padx=5)
        ttk.Button(btns_frame, text="Bazani filtrlash", style='Accent.TButton', 
                   command=lambda: self.show_admin_baza_filter(admin_win)).pack(side='left', padx=5)
        ttk.Button(btns_frame, text="O'quvchilar ro'yxati", style='Accent.TButton', 
                   command=self.show_students_list).pack(side='left', padx=5)
        ttk.Button(btns_frame, text="Statistika", style='Accent.TButton',
                   command=lambda: self.show_statistics(admin_win)).pack(side='left', padx=5)
        # Yangi "Bosh menyu" tugmasi
        ttk.Button(btns_frame, text="Bosh menyu", style='Warning.TButton',
                   command=self.return_to_main_menu).pack(side='left', padx=5)

        # --- SAVOL TAHRIRLASH BO'LIMI (doimiy ko'rinadi) ---
        # Fayl tanlash
        files = glob.glob(os.path.join(BASE_DIR, "baza", "*.docx"))
        file_names = [os.path.basename(f) for f in files]
        file_var = tk.StringVar(value=file_names[0] if file_names else "")
        file_combo = ttk.Combobox(admin_win, values=file_names, textvariable=file_var, state='readonly', font=('Segoe UI', 12))
        file_combo.pack(pady=5)
        
        # Index entry o'zgarishini kuzatish uchun yangi funksiya
        def on_index_change(*args):
            idx = index_var.get().strip()
            if idx.isdigit() and int(idx) > 0:
                self.admin_show_question(file_var.get(), question_entry, variant_entries, correct_var, admin_win)
        
        # Savol raqamini tanlash (tahrirlash uchun)
        ttk.Label(admin_win, text="Savol raqami (tahrirlash uchun, bo'sh qoldirsangiz yangi savol qo'shiladi):", font=('Segoe UI', 11)).pack(pady=(10, 0))
        index_var = tk.StringVar()
        index_var.trace('w', on_index_change)  # 'write' o'rniga 'w' ishlatamiz
        index_entry = ttk.Entry(admin_win, font=('Segoe UI', 12), textvariable=index_var)
        index_entry.pack(pady=5)
        
        # Savol va variantlar
        ttk.Label(admin_win, text="Savol matni:", font=('Segoe UI', 11)).pack(pady=(10, 0))
        question_entry = tk.Text(admin_win, height=3, font=('Segoe UI', 12), wrap='word')
        question_entry.pack(pady=5, fill='x', padx=10)
        variant_entries = []
        for i in range(4):
            frame = ttk.Frame(admin_win)
            frame.pack(fill='x', padx=10, pady=2)
            ttk.Label(frame, text=f"Variant {chr(65+i)}:", font=('Segoe UI', 11)).pack(side='left')
            entry = ttk.Entry(frame, font=('Segoe UI', 12), width=60)
            entry.pack(side='left', padx=5)
            variant_entries.append(entry)
        # To'g'ri javob tanlash
        correct_var = tk.IntVar(value=0)
        correct_frame = ttk.Frame(admin_win)
        correct_frame.pack(pady=5)
        ttk.Label(correct_frame, text="To'g'ri javob:", font=('Segoe UI', 11)).pack(side='left')
        for i in range(4):
            ttk.Radiobutton(correct_frame, text=chr(65+i), variable=correct_var, value=i).pack(side='left', padx=5)
        # Rasm bilan ishlash uchun o'zgaruvchilar
        def add_image():
            file_path = filedialog.askopenfilename(filetypes=[("Image files", "*.png;*.jpg;*.jpeg;*.bmp")])
            if file_path:
                self.admin_image_path = file_path
                img = Image.open(file_path)
                img.thumbnail((120, 120))
                self.admin_image_tk = ImageTk.PhotoImage(img)
                image_label.config(image=self.admin_image_tk)
                image_label.image = self.admin_image_tk
        def paste_image():
            try:
                img = ImageGrab.grabclipboard()
                if isinstance(img, Image.Image):
                    temp_path = os.path.join(tempfile.gettempdir(), f"admin_paste_{random.randint(1000,9999)}.png")
                    img.save(temp_path)
                    self.admin_image_path = temp_path
                    img.thumbnail((120, 120))
                    self.admin_image_tk = ImageTk.PhotoImage(img)
                    image_label.config(image=self.admin_image_tk)
                    image_label.image = self.admin_image_tk
                else:
                    messagebox.showerror("Xato", "Clipboard'da rasm topilmadi!")
            except Exception as e:
                messagebox.showerror("Xato", f"Clipboard'dan rasm olishda xato: {e}")
        def remove_image():
            self.admin_image_path = None
            image_label.config(image='')
            image_label.image = None
        # Rasm tugmalari va label
        img_btn_frame = ttk.Frame(admin_win)
        img_btn_frame.pack(pady=5)
        ttk.Button(img_btn_frame, text="Rasm qo'shish", command=add_image).pack(side='left', padx=5)
        ttk.Button(img_btn_frame, text="Paste rasm", command=paste_image).pack(side='left', padx=5)
        ttk.Button(img_btn_frame, text="Rasmni olib tashlash", command=remove_image).pack(side='left', padx=5)
        image_label = ttk.Label(admin_win)
        image_label.pack(pady=5)
        
        # Ko'rish tugmasini olib tashlab, faqat asosiy tugmalarni qo'yamiz
        action_btn_frame = ttk.Frame(admin_win)
        action_btn_frame.pack(pady=10)
        ttk.Button(action_btn_frame, text="Yangi savol qo'shish", style='Primary.TButton', 
                   command=lambda: self.admin_add_question(file_var.get(), question_entry, variant_entries, correct_var, admin_win)).pack(side='left', padx=5)
        ttk.Button(action_btn_frame, text="Tahrirlash (saqlash)", style='Accent.TButton', 
                   command=lambda: self.admin_edit_question(file_var.get(), index_entry, question_entry, variant_entries, correct_var, admin_win)).pack(side='left', padx=5)
        ttk.Button(action_btn_frame, text="Yopish", style='Danger.TButton', 
                   command=admin_win.destroy).pack(side='left', padx=5)

    def show_admin_settings(self, parent):
        win = Toplevel(parent)
        win.title("Sozlamalar")
        win.geometry("500x300")
        win.grab_set()
        ttk.Label(win, text="Test vaqti (daqiqa):", font=('Segoe UI', 11)).grid(row=0, column=0, sticky='w')
        time_entry = ttk.Entry(win, font=('Segoe UI', 11), width=10)
        time_entry.insert(0, str(self.settings["test_time"]))
        time_entry.grid(row=0, column=1, padx=5)
        ttk.Label(win, text="Savollar soni:", font=('Segoe UI', 11)).grid(row=0, column=2, sticky='w')
        questions_entry = ttk.Entry(win, font=('Segoe UI', 11), width=10)
        questions_entry.insert(0, str(self.settings["questions_per_test"]))
        questions_entry.grid(row=0, column=3, padx=5)
        shuffle_var = tk.BooleanVar(value=self.settings["shuffle_answers"])
        ttk.Checkbutton(win, text="Javoblarni aralashtirish", variable=shuffle_var).grid(row=1, column=0, sticky='w', pady=5)
        dark_mode_var = tk.BooleanVar(value=self.settings["dark_mode"])
        ttk.Checkbutton(win, text="Qorong'i rejim", variable=dark_mode_var, command=lambda: self.toggle_dark_mode(dark_mode_var.get())).grid(row=1, column=1, sticky='w', pady=5)
        ttk.Label(win, text="Admin parolini o'zgartirish:", font=('Segoe UI', 11)).grid(row=2, column=0, sticky='w', pady=(10,0))
        current_pass = ttk.Entry(win, show='*', width=15)
        current_pass.grid(row=2, column=1, padx=5, pady=(10,0))
        new_pass = ttk.Entry(win, show='*', width=15)
        new_pass.grid(row=2, column=3, padx=5, pady=(10,0))
        ttk.Label(win, text="Joriy parol:").grid(row=2, column=0, sticky='e', pady=(10,0))
        ttk.Label(win, text="Yangi parol:").grid(row=2, column=2, sticky='e', pady=(10,0))
        def save_settings():
            password = simpledialog.askstring(
                "Admin paroli",
                "Sozlamalarni saqlash uchun admin parolini kiriting:",
                show='*',
                parent=win
            )
            if password != self.admin_password:
                messagebox.showerror("Xato", "Noto'g'ri admin paroli! Sozlamalar saqlanmadi.")
                return
            try:
                new_time = int(time_entry.get())
                new_questions = int(questions_entry.get())
                if new_time <= 0 or new_questions <= 0:
                    raise ValueError("Qiymatlar musbat son bo'lishi kerak")
                self.settings.update({
                    "test_time": new_time,
                    "questions_per_test": new_questions,
                    "shuffle_answers": shuffle_var.get(),
                    "dark_mode": dark_mode_var.get()
                })
                if current_pass.get() and new_pass.get():
                    if current_pass.get() == self.admin_password:
                        self.admin_password = new_pass.get()
                        messagebox.showinfo("Muvaffaqiyat", "Parol muvaffaqiyatli o'zgartirildi!")
                    else:
                        messagebox.showerror("Xato", "Noto'g'ri joriy parol!")
                self.save_settings()
                messagebox.showinfo("Muvaffaqiyat", "Sozlamalar saqlandi!")
            except ValueError as e:
                messagebox.showerror("Xato", f"Noto'g'ri qiymat: {e}")
        ttk.Button(win, text="Sozlamalarni saqlash", style='Primary.TButton', command=save_settings).grid(row=3, column=0, columnspan=4, pady=10)
        ttk.Button(win, text="Yopish", style='Danger.TButton', command=win.destroy).grid(row=4, column=0, columnspan=4, pady=10)

    def show_admin_baza_filter(self, parent):
        win = Toplevel(parent)
        win.title("Bazani filtrlash")
        win.geometry("400x500")
        win.grab_set()
        files = glob.glob(os.path.join(BASE_DIR, "baza", "*.docx"))
        file_names = [os.path.basename(f) for f in files]
        enabled_files = set()
        if os.path.exists('enabled_baza_files.json'):
            try:
                with open('enabled_baza_files.json', 'r') as f:
                    enabled_files = set(json.load(f))
            except Exception:
                enabled_files = set()
        file_vars = {}
        for i, fname in enumerate(file_names):
            var = tk.BooleanVar(value=(fname in enabled_files or not enabled_files))
            cb = ttk.Checkbutton(win, text=fname, variable=var)
            cb.grid(row=i, column=0, sticky='w')
            file_vars[fname] = var
        def save_enabled_files():
            selected = [fname for fname, var in file_vars.items() if var.get()]
            with open('enabled_baza_files.json', 'w') as f:
                json.dump(selected, f)
            messagebox.showinfo("Admin", "Tanlangan bazalar saqlandi!", parent=win)
        ttk.Button(win, text="Bazalarni saqlash", style='Primary.TButton', command=save_enabled_files).grid(row=len(file_names)+1, column=0, pady=10)
        ttk.Button(win, text="Yopish", style='Danger.TButton', command=win.destroy).grid(row=len(file_names)+2, column=0, pady=10)

    def admin_add_question(self, file_name, question_entry, variant_entries, correct_var, win):
        from docx import Document
        import os
        file_path = os.path.join(BASE_DIR, "baza", file_name)
        try:
            doc = Document(file_path)
            table = doc.tables[0]  # Faraz: 1-jadval test jadvali
            new_row = table.add_row()
            cells = new_row.cells
            # Savol matni
            cells[0].text = question_entry.get("1.0", "end").strip()
            # Rasmni savolga joylashtirish
            if self.admin_image_path:
                paragraph = cells[0].paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(self.admin_image_path, width=docx.shared.Inches(1.5))
            # To'g'ri javob
            correct_idx = correct_var.get()
            cells[1].text = variant_entries[correct_idx].get().strip()
            # Noto'g'ri javoblar
            false_indices = [i for i in range(4) if i != correct_idx]
            for j, idx in enumerate(false_indices):
                cells[2+j].text = variant_entries[idx].get().strip()
            doc.save(file_path)
            messagebox.showinfo("Admin", f"Yangi savol '{file_name}' fayliga muvaffaqiyatli qo'shildi!")
        except Exception as e:
            messagebox.showerror("Xato", f"Savol qo'shishda xato: {e}")

    def admin_edit_question(self, file_name, index_entry, question_entry, variant_entries, correct_var, admin_win):
        from docx import Document
        import os
        file_path = os.path.join(BASE_DIR, "baza", file_name)
        try:
            doc = Document(file_path)
            table = doc.tables[0]
            idx = index_entry.get().strip()
            if not idx.isdigit() or int(idx) < 1 or int(idx) > len(table.rows)-1:
                messagebox.showerror("Xato", "To'g'ri savol raqamini kiriting!")
                return
            row_idx = int(idx)
            row = table.rows[row_idx]
            cells = row.cells
            cells[0].text = question_entry.get("1.0", "end").strip()
            # Rasmni savolga joylashtirish
            if self.admin_image_path:
                paragraph = cells[0].paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(self.admin_image_path, width=docx.shared.Inches(1.5))
            correct_idx = correct_var.get()
            cells[1].text = variant_entries[correct_idx].get().strip()
            false_indices = [i for i in range(4) if i != correct_idx]
            for j, idx2 in enumerate(false_indices):
                cells[2+j].text = variant_entries[idx2].get().strip()
            doc.save(file_path)
            messagebox.showinfo("Admin", f"{file_name} faylida savol muvaffiyatli tahrirlandi!")
            # admin_win.destroy() yo'q!
        except Exception as e:
            messagebox.showerror("Xato", f"Savolni tahrirlashda xato: {e}")

    def admin_show_question(self, file_name, question_entry, variant_entries, correct_var, win):
        from docx import Document
        import os
        file_path = os.path.join(BASE_DIR, "baza", file_name)
        try:
            doc = Document(file_path)
            table = doc.tables[0]
            # Indexni admin panelidagi index_entry maydonidan olamiz
            idx_entry = None
            for widget in win.winfo_children():
                if isinstance(widget, ttk.Entry) and widget.get().isdigit():
                    idx_entry = widget
                    break
            if idx_entry is None:
                # Yoki, index_entry ni argument sifatida uzatish kerak bo'lishi mumkin
                idx_entry = win.nametowidget('index_entry') if 'index_entry' in win.children else None
            idx = None
            if idx_entry:
                idx = idx_entry.get().strip()
            if not idx or not idx.isdigit() or int(idx) < 1 or int(idx) > len(table.rows) - 1:
                messagebox.showerror("Xato", "To'g'ri savol raqamini kiriting!")
                return
            row = table.rows[int(idx)]
            cells = row.cells
            if len(cells) < 5:
                messagebox.showerror("Xato", "Bu savolda kamida 5 ta ustun (cell) bo'lishi kerak!")
                return
            # Savol matni
            question_entry.delete("1.0", "end")
            question_entry.insert("1.0", cells[0].text.strip())
            # Variantlar
            for i in range(4):
                variant_entries[i].delete(0, "end")
            # To'g'ri javob (always 1st variant in docx)
            variant_entries[0].insert(0, cells[1].text.strip())
            # Noto'g'ri javoblar
            for j in range(3):
                variant_entries[j+1].insert(0, cells[2+j].text.strip())
            # To'g'ri javobni tanlash (always 0)
            correct_var.set(0)
            # messagebox.showinfo("Ko'rish", f"{idx}-savol ko'rsatildi!")
        except Exception as e:
            messagebox.showerror("Xato", f"Savolni ko'rsatishda xato: {e}")

    def show_students_list(self):
        win = Toplevel(self.root)
        win.title("O'quvchilar ro'yxati")
        win.geometry("700x500")
        win.grab_set()
        frame = ttk.Frame(win, padding=20)
        frame.pack(expand=True, fill='both')
        listbox = tk.Listbox(frame, font=('Segoe UI', 12), selectmode='single')
        listbox.pack(fill='both', expand=True, pady=10)
        ism_entry = ttk.Entry(frame, font=('Segoe UI', 12), width=15)
        ism_entry.pack(fill='x', pady=2)
        familiya_entry = ttk.Entry(frame, font=('Segoe UI', 12), width=15)
        familiya_entry.pack(fill='x', pady=2)
        login_entry = ttk.Entry(frame, font=('Segoe UI', 12), width=15)
        login_entry.pack(fill='x', pady=2)
        parol_entry = ttk.Entry(frame, font=('Segoe UI', 12), width=15)
        parol_entry.pack(fill='x', pady=2)
        def refresh_listbox():
            listbox.delete(0, 'end')
            for s in read_students_list():
                listbox.insert('end', f"{s['ism']} | {s['familiya']} | {s['login']} | {s['parol']}")
        def add_student():
            ism = ism_entry.get().strip().upper()
            familiya = familiya_entry.get().strip().upper()
            login = login_entry.get().strip()
            parol = parol_entry.get().strip()
            if ism and familiya and login and parol and not any(s['login'] == login for s in read_students_list()):
                add_student_to_list(ism, familiya, login, parol)
                refresh_listbox()
                ism_entry.delete(0, 'end')
                familiya_entry.delete(0, 'end')
                login_entry.delete(0, 'end')
                parol_entry.delete(0, 'end')
        def remove_student():
            sel = listbox.curselection()
            if sel:
                login = listbox.get(sel[0]).split('|')[2].strip()
                remove_student_from_list(login)
                refresh_listbox()
        btn_frame = ttk.Frame(frame)
        btn_frame.pack(pady=5)
        ttk.Button(btn_frame, text="Qo'shish", style='Primary.TButton', command=add_student).pack(side='left', padx=5)
        ttk.Button(btn_frame, text="O'chirish", style='Danger.TButton', command=remove_student).pack(side='left', padx=5)
        ttk.Button(btn_frame, text="Yopish", style='Accent.TButton', command=win.destroy).pack(side='left', padx=5)
        refresh_listbox()

    def save_test_result(self, student_name, score, total_questions):
        if not student_name or score < 0 or total_questions <= 0:
            print("Noto'g'ri test ma'lumotlari")
            return
            
        try:
            # Test natijasini API ga yuborish
            answers = {
                str(i): {
                    'question_id': self.current_test_set[i].get('id', ''),
                    'selected_answer': self.user_answers.get(i),
                    'correct_answer': self.current_test_set[i].get('correct', 0),
                    'is_correct': self.user_answers.get(i) == self.current_test_set[i].get('correct', 0)
                }
                for i in range(len(self.current_test_set))
            }
            
            api_success = api_save_result(student_name, score, total_questions, answers)
            
            if not api_success:
                print("API ga natijani yuborishda xatolik")
                # Lokal faylga zaxira sifatida saqlash
                self.save_result_locally(student_name, score, total_questions)
                
        except Exception as e:
            print(f"Natijani saqlashda xato: {str(e)}")
            # Xatolik yuz berganda lokal faylga saqlash
            self.save_result_locally(student_name, score, total_questions)
            
    def save_result_locally(self, student_name, score, total_questions):
        """Natijani lokal faylga saqlash"""
        try:
            results_file = 'test_results.json'
            current_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            
            # Mavjud natijalarni o'qish
            results = []
            if os.path.exists(results_file):
                try:
                    with open(results_file, 'r', encoding='utf-8') as f:
                        results = json.load(f)
                    if not isinstance(results, list):
                        results = []
                except json.JSONDecodeError:
                    results = []
            
            # Yangi natijani qo'shish
            new_result = {
                'student_name': student_name.strip(),
                'score': score,
                'total_questions': total_questions,
                'percentage': round((score / total_questions) * 100, 2),
                'date': current_time
            }
            results.append(new_result)
            
            # Natijalarni saqlash
            with open(results_file, 'w', encoding='utf-8') as f:
                json.dump(results, f, ensure_ascii=False, indent=4)
                
        except Exception as e:
            print(f"Lokal faylga saqlashda xato: {str(e)}")
            try:
                # Backup fayl yaratish
                backup_file = f'test_results_backup_{int(time.time())}.json'
                with open(backup_file, 'w', encoding='utf-8') as f:
                    json.dump([new_result], f, ensure_ascii=False, indent=4)
                print(f"Natija backup faylga saqlandi: {backup_file}")
            except Exception as be:
                print(f"Backup yaratishda xato: {str(be)}")

    def get_test_results(self):
        try:
            results_file = 'test_results.json'
            if not os.path.exists(results_file):
                return []

            try:
                with open(results_file, 'r', encoding='utf-8') as f:
                    all_results = json.load(f)
                if not isinstance(all_results, list):
                    return []
            except json.JSONDecodeError:
                return []

            # Natijalarni o'quvchilar bo'yicha guruhlash
            student_results = {}
            for result in all_results:
                if not isinstance(result, dict) or 'student_name' not in result:
                    continue
                    
                student_name = result.get('student_name', '').strip()
                if not student_name:
                    continue
                    
                if student_name not in student_results:
                    student_results[student_name] = {
                        'total_tests': 0,
                        'total_score': 0,
                        'scores': [],
                        'last_test': result.get('date', '')
                    }
                
                sr = student_results[student_name]
                sr['total_tests'] += 1
                percentage = result.get('percentage', 0)
                sr['total_score'] += percentage
                sr['scores'].append(percentage)
                
                result_date = result.get('date', '')
                if result_date and result_date > sr['last_test']:
                    sr['last_test'] = result_date

            # Natijalarni formatlash
            formatted_results = []
            for student_name, data in student_results.items():
                if data['total_tests'] > 0:  # Nolga bo'lishni oldini olish
                    formatted_results.append({
                        'student_name': student_name,
                        'total_tests': data['total_tests'],
                        'avg_score': round(data['total_score'] / data['total_tests'], 2),
                        'last_test': data['last_test'],
                        'best_score': max(data['scores']) if data['scores'] else 0
                    })

            # O'rtacha ball bo'yicha saralash
            formatted_results.sort(key=lambda x: x['avg_score'], reverse=True)
            return formatted_results

        except Exception as e:
            print(f"Natijalarni olishda xato: {str(e)}")
            return []

    def show_statistics(self, parent):
        if not parent or not parent.winfo_exists():
            return
            
        try:
            stats_win = Toplevel(parent)
            stats_win.title("Test Statistikasi")
            stats_win.geometry("900x700")
            stats_win.grab_set()
            
            # Asosiy frame
            main_frame = ttk.Frame(stats_win, padding=20)
            main_frame.pack(expand=True, fill='both')

            # Umumiy statistika
            general_stats = ttk.LabelFrame(main_frame, text="Umumiy ma'lumotlar", padding=15)
            general_stats.pack(fill='x', pady=(0, 20))

            try:
                # O'quvchilar statistikasi
                students = read_students_list()
                ttk.Label(general_stats, 
                         text=f"Ro'yxatdagi o'quvchilar soni: {len(students)}", 
                         font=('Segoe UI', 12)).pack(anchor='w', pady=2)

                # Test bazasi statistikasi
                baza_path = os.path.join(BASE_DIR, "baza")
                if not os.path.exists(baza_path):
                    os.makedirs(baza_path, exist_ok=True)
                    messagebox.showwarning("Ogohlantirish", "Test bazasi papkasi mavjud emas. Yangi papka yaratildi.")

                files = glob.glob(os.path.join(baza_path, "*.docx"))
                total_questions = 0
                valid_files = 0

                for file in files:
                    try:
                        doc = Document(file)
                        if doc.tables and len(doc.tables) > 0:
                            rows = len(doc.tables[0].rows) - 1
                            if rows > 0:  # Faqat haqiqiy savollar sonini qo'shish
                                total_questions += rows
                                valid_files += 1
                    except Exception as e:
                        print(f"Faylni o'qishda xato {file}: {str(e)}")
                        continue

                ttk.Label(general_stats, 
                         text=f"Test fayllari soni: {valid_files}", 
                         font=('Segoe UI', 12)).pack(anchor='w', pady=2)
                ttk.Label(general_stats, 
                         text=f"Jami savollar soni: {total_questions}", 
                         font=('Segoe UI', 12)).pack(anchor='w', pady=2)

            except Exception as e:
                messagebox.showerror("Xato", f"Statistika ma'lumotlarini olishda xato: {str(e)}")
                return

            # Test natijalari jadvali
            results_frame = ttk.LabelFrame(main_frame, text="Test natijalari", padding=15)
            results_frame.pack(fill='both', expand=True, pady=(0, 20))

            # Ustunlar
            columns = ('rank', 'name', 'tests', 'avg_score', 'last_test', 'best_score')
            tree = ttk.Treeview(results_frame, columns=columns, show='headings')

            # Ustun sozlamalari
            tree.heading('rank', text='№')
            tree.heading('name', text='F.I.O')
            tree.heading('tests', text='Testlar soni')
            tree.heading('avg_score', text="O'rtacha ball")
            tree.heading('last_test', text='Oxirgi test')
            tree.heading('best_score', text='Eng yuqori ball')

            # Ustun o'lchamlari
            tree.column('rank', width=50, anchor='center')
            tree.column('name', width=200)
            tree.column('tests', width=100, anchor='center')
            tree.column('avg_score', width=100, anchor='center')
            tree.column('last_test', width=150, anchor='center')
            tree.column('best_score', width=100, anchor='center')

            # Scrollbar
            scrollbar = ttk.Scrollbar(results_frame, orient='vertical', command=tree.yview)
            tree.configure(yscrollcommand=scrollbar.set)
            
            tree.pack(side='left', fill='both', expand=True)
            scrollbar.pack(side='right', fill='y')

            def refresh_results():
                try:
                    # Jadvalni tozalash
                    for item in tree.get_children():
                        tree.delete(item)

                    # Test natijalarini o'qish
                    results = self.get_test_results()
                    
                    # Natijalarni saralash va ko'rsatish
                    for i, result in enumerate(results, 1):
                        tree.insert('', 'end', values=(
                            i,
                            result['student_name'],
                            result['total_tests'],
                            f"{result['avg_score']}%",
                            result['last_test'],
                            f"{result['best_score']}%"
                        ))
                except Exception as e:
                    messagebox.showerror("Xato", f"Natijalarni yangilashda xato: {str(e)}")

            # Boshqaruv tugmalari
            control_frame = ttk.Frame(main_frame)
            control_frame.pack(fill='x', pady=10)

            def export_excel():
                try:
                    filename = filedialog.asksaveasfilename(
                        defaultextension=".xlsx",
                        filetypes=[("Excel files", "*.xlsx")],
                        initialfile=f"test_statistika_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx"
                    )
                    if filename:
                        wb = openpyxl.Workbook()
                        ws = wb.active
                        ws.title = "Test Statistikasi"

                        # Sarlavhalar
                        headers = ['№', 'F.I.O', 'Testlar soni', "O'rtacha ball", 
                                 'Oxirgi test', 'Eng yuqori ball']
                        ws.append(headers)

                        # Ma'lumotlar
                        for item in tree.get_children():
                            ws.append(tree.item(item)['values'])

                        # Formatlash
                        for col in ws.columns:
                            max_length = 0
                            for cell in col:
                                try:
                                    max_length = max(max_length, len(str(cell.value)))
                                except:
                                    pass
                            adjusted_width = (max_length + 2)
                            ws.column_dimensions[openpyxl.utils.get_column_letter(col[0].column)].width = adjusted_width

                        wb.save(filename)
                        messagebox.showinfo("Muvaffaqiyat", "Statistika muvaffaqiyatli eksport qilindi!")
                except Exception as e:
                    messagebox.showerror("Xato", f"Eksport qilishda xato: {str(e)}")

            ttk.Button(control_frame, text="Excel'ga eksport", 
                      style='Primary.TButton', command=export_excel).pack(side='left', padx=5)
            ttk.Button(control_frame, text="Yangilash", 
                      style='Accent.TButton', command=refresh_results).pack(side='left', padx=5)
            ttk.Button(control_frame, text="Yopish", 
                      style='Danger.TButton', command=stats_win.destroy).pack(side='right', padx=5)

            # Dastlabki ma'lumotlarni yuklash
            refresh_results()

        except Exception as e:
            messagebox.showerror("Xato", f"Statistika oynasini ochishda xato: {str(e)}")

    def return_to_main_menu(self):
        # Foydalanuvchidan tasdiqlash so'rash
        if not messagebox.askyesno("Tasdiqlash", 
            "Bosh menyuga qaytmoqchimisiz?\nBarcha joriy ma'lumotlar o'chiriladi."):
            return
            
        try:
            # Agar test yakunlangan bo'lsa va xatolar ko'rib chiqilmagan bo'lsa
            if hasattr(self, 'review_start_time') and self.review_start_time:
                self.send_review_status(force_send=True)
                
            # Timer va boshqa jarayonlarni to'xtatish
            self.timer_running = False
            self.timer_seconds = 0
            
            # O'zgaruvchilarni tozalash
            self.username = ""
            self.all_questions = []
            self.current_test_set = []
            self.user_answers = {}
            self.current_question = 0
            self.results_shown = False
            self.admin_override = False
            self.current_zoom = 1.0
            self.used_question_indices = set()
            
            # Barcha oynalarni yopish
            for widget in self.root.winfo_children():
                if isinstance(widget, Toplevel):
                    widget.destroy()
            for window in self.root.winfo_toplevel().winfo_children():
                if isinstance(window, Toplevel):
                    window.destroy()
            
            # Bosh menyuni qayta yaratish
            self.create_start_screen()
            
        except Exception as e:
            messagebox.showerror("Xato", f"Bosh menyuga qaytishda xatolik: {str(e)}")
            # Xato yuz berganda ham bosh menyuni ko'rsatishga harakat qilish
            try:
                self.create_start_screen()
            except:
                pass

    def send_review_status(self, force_send=False):
        """Test yakunlangandan so'ng xatolarni tahlil qilish holati haqida ma'lumot yuborish"""
        try:
            if not hasattr(self, 'wrong_answers') or not hasattr(self, 'review_start_time') or not self.review_start_time:
                print("Tahlil ma'lumotlari mavjud emas")
                return
                
            total_wrong = len(self.wrong_answers)
            if total_wrong == 0:  # Agar xatolar yo'q bo'lsa
                print("Xatolar yo'q, tahlil yuborilmaydi")
                return
                
            reviewed = len(self.reviewed_questions)
            review_duration = datetime.now() - self.review_start_time
            
            # Agar majburiy yuborish bo'lmasa va tahlil vaqti 5 soniyadan kam bo'lsa, hisobga olmaslik
            if not force_send and review_duration.seconds < 5:
                print("Tahlil vaqti juda qisqa, yuborilmaydi")
                return
                
            status_text = (
                f"🔍 <b>Xatolar tahlili:</b>\n"
                f"<b>O'quvchi:</b> {self.username}\n"
                f"<b>Jami xatolar:</b> {total_wrong}\n"
                f"<b>Ko'rib chiqilgan xatolar:</b> {reviewed}\n"
                f"<b>Ko'rib chiqilmagan xatolar:</b> {total_wrong - reviewed}\n"
                f"<b>Tahlil vaqti:</b> {review_duration.seconds//60} daqiqa {review_duration.seconds%60} soniya\n\n"
            )
            
            if reviewed == 0:
                status_text += "⚠️ <b>O'quvchi xatolarini umuman ko'rib chiqmadi!</b>"
            elif reviewed < total_wrong:
                status_text += f"⚠️ <b>O'quvchi {total_wrong - reviewed} ta xatosini ko'rib chiqmadi!</b>"
            else:
                status_text += "✅ <b>Barcha xatolar ko'rib chiqildi!</b>"
                
            print(f"\nTelegramga yuborilayotgan xabar:\n{status_text}")
                
            # Telegramga yuborish
            chat_id = None
            if os.path.exists('telegram_chat_id.txt'):
                try:
                    with open('telegram_chat_id.txt', 'r') as f:
                        chat_id = f.read().strip()
                        print(f"Saqlangan chat ID o'qildi: {chat_id}")
                except Exception as e:
                    print(f"Chat ID faylini o'qishda xato: {str(e)}")
            
            global global_admin_chat_id
            if not chat_id and global_admin_chat_id:
                chat_id = global_admin_chat_id
                print(f"Global chat ID ishlatilmoqda: {chat_id}")
                
            if not chat_id:
                print("Chat ID topilmadi, yangi chat ID so'ralmoqda...")
                chat_id = get_latest_chat_id(TELEGRAM_BOT_TOKEN)
                if chat_id:
                    try:
                        with open('telegram_chat_id.txt', 'w') as f:
                            f.write(chat_id)
                        global_admin_chat_id = chat_id
                        print(f"Yangi chat ID saqlandi: {chat_id}")
                    except Exception as e:
                        print(f"Chat ID ni saqlashda xato: {str(e)}")
                    
            if chat_id:
                print(f"Xabar {chat_id} ID'li chatga yuborilmoqda...")
                success = send_result_to_telegram(status_text, TELEGRAM_BOT_TOKEN, chat_id)
                if success:
                    print("Xabar muvaffaqiyatli yuborildi!")
                else:
                    print("Xabar yuborishda xatolik yuz berdi!")
            else:
                print("Telegram chat_id topilmadi!")
                
            # O'zgaruvchilarni tozalash (faqat majburiy yuborish bo'lmaganda)
            if not force_send:
                self.review_start_time = None
                self.reviewed_questions.clear()
                if hasattr(self, 'wrong_answers'):
                    del self.wrong_answers
                
        except Exception as e:
            print(f"Tahlil holatini yuborishda xatolik: {str(e)}")
            traceback.print_exc()

if __name__ == "__main__":
    root = tk.Tk()
    app = ModernTestApp(root)
    root.mainloop()